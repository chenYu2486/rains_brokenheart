# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2

from docx import Document


BASE = Path(r"D:\rains_brokenheart-main\rains_brokenheart-main\Graduation Thesis")
SOURCE = max(
    [
        p for p in BASE.glob("*.docx")
        if "实验图表与量化评测增强版" in p.name and not p.name.startswith("~$")
    ],
    key=lambda p: p.stat().st_mtime,
)
OUTPUT = BASE / f"{SOURCE.stem}_综合增强版.docx"


def replace_startswith(doc, prefix, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = text
            return True
    return False


def replace_contains(doc, needle, text):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = text
            return True
    return False


def main():
    copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    replace_startswith(
        doc,
        "摘要：",
        "摘要：针对心理健康服务可及性不足、传统对话系统缺乏状态跟踪与专业边界约束等问题，本文设计并实现了“避雨檐 ShelterAI”精神心理对话系统。系统采用评估模型与陪伴模型分阶段协作架构：前者负责建档访谈、周期复评和结构化JSON画像生成，后者基于最新画像开展支持性对话。为提升领域适配能力，本文构建约2万条心理健康多轮指令数据，并基于LoRA开展多组对照实验。实验图表显示，Proposed LoRA-Optimized方案在评估损失、准确率和训练—评估损失差距之间取得较好平衡。系统还集成本地RAG知识库、长期记忆、Visual Novel界面和云端代理转发机制。测试表明，原型能够完成关键词定制、初评建档、画像更新、知识来源回显和异常提示等核心流程，具有一定工程复现价值。",
    )
    replace_startswith(
        doc,
        "ABSTRACT:",
        "ABSTRACT: This thesis presents ShelterAI, a mental health dialogue prototype based on large language model fine-tuning. It adopts a staged dual-model architecture: an assessment model builds structured JSON profiles, while a LoRA-adapted companion model provides supportive dialogue. About 20,000 multi-turn instruction samples are constructed, and comparative LoRA experiments are reported with loss, accuracy, gap, gradient, and learning-rate curves. The system also integrates a local RAG knowledge base, long-term memory, a Visual Novel interface, and cloud proxy forwarding. Tests verify intake assessment, profile updating, citation display, and fault handling.",
    )

    replace_contains(
        doc,
        "五是结合LoRA微调方案讨论后续接入微调模型的可行性",
        "围绕精神心理对话系统的工程实现，本文主要完成五项工作：一是建立面向心理健康初筛的需求模型，将系统功能划分为关键词定制、建档访谈、结构化评估、持续陪伴、知识库检索、风险提示和会话管理等模块；二是构建评估模型与陪伴模型分阶段协作机制，将建档任务与陪伴任务解耦；三是实现本地RAG检索模块，把DSM-5-TR相关资料转换为带页码、哈希和片段编号的知识单元；四是完成浏览器端流式对话界面、画像面板和存档管理；五是整理LoRA训练实验结果，将微调模型以qwen3-8b-9c3af956383a的模型ID接入陪伴模型链路。",
    )

    replace_startswith(
        doc,
        "LoRA微调实验基于Qwen3.5-9B基座模型",
        "LoRA微调实验基于Qwen系列中文对话基座模型展开，训练完成后合并权重并部署为qwen3-8b-9c3af956383a。实验采用LoRA对注意力层的Q和K投影矩阵进行低秩适配，使用自建心理健康对话数据集，共约20,000条多轮对话样本，覆盖20个心理话题标签（如职场内耗、原生家庭、亲密关系等），每条样本包含system、user、assistant多轮交互。数据集按照8:1:1的比例划分为训练集、验证集和测试集。",
    )

    replace_startswith(
        doc,
        "为进一步展示训练过程，本文将训练损失",
        "为进一步展示训练过程，本文将训练损失、评估损失、评估准确率、训练—评估损失差距、梯度范数和学习率调度分别绘制为图5-1、图5-2、图5-3、图5-4、图5-5和图5-6。其中，图5-1反映模型对训练数据的拟合速度，图5-2和图5-3反映验证集上的泛化表现，图5-4用于比较过拟合风险，图5-5用于观察梯度稳定性，图5-6用于说明不同学习率策略的调度过程。",
    )

    replace_startswith(
        doc,
        "一是当前项目仓库主要体现前端原型和本地RAG链路",
        "一是当前微调实验虽然已经形成训练曲线、评估曲线和对照结果，但评价仍以自动指标和工程流程验证为主，尚缺少由心理咨询专业人员参与的人工评分与安全审查。因此，后续应进一步补充人工评测、风险场景测试和跨数据集泛化验证。",
    )

    replace_startswith(
        doc,
        "项目源码表明，系统已实现关键词定制",
        "项目源码和实验结果表明，系统已实现关键词定制、建档访谈、评估画像、持续陪伴、周期复评、本地知识库检索、流式API、会话存档和模型阶段可视化等核心功能。RAG模块将心理健康相关资料转换为带页码和哈希的知识片段，并在检索命中后注入提示词，使专业性回答具有一定溯源依据；LoRA实验则从训练损失、评估损失、准确率和泛化差距等角度验证了微调方案的有效性。",
    )

    replace_startswith(
        doc,
        "后续工作可从四个方面展开",
        "后续工作可从四个方面展开：一是扩大微调实验与自动化评测规模，加入人工抽检、安全红队测试、不同基座模型对比和消融实验；二是升级RAG检索能力，引入语义嵌入模型、向量数据库和重排序器，提高对中文口语化表达的召回能力；三是强化安全机制，在输入端、生成端和输出端分别加入风险识别、危机模板、人工转介和输出审核；四是完善隐私与部署方案，加入本地加密、数据脱敏、权限管理和用户知情同意机制。",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
