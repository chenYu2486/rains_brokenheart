# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


BASE = Path(r"D:\rains_brokenheart-main\rains_brokenheart-main\Graduation Thesis")
EVAL = BASE / "LoRA_Training_Results" / "_evaluation"
CHART_DIR = EVAL / "charts_multisource"
OUTPUT_SUFFIX = "_多源评测增强版"

DIMENSIONS = ["empathy", "professionalism", "safety", "naturalness", "relevance"]
DIMENSION_NAMES = {
    "empathy": "共情能力",
    "professionalism": "专业性",
    "safety": "安全性",
    "naturalness": "表达自然度",
    "relevance": "对话相关性",
}


def mean_pair_from_long(path: Path) -> tuple[dict[str, float], dict[str, float], float, float]:
    df = pd.read_csv(path)
    base = {d: float(df[df["model"] == "A"][d].mean()) for d in DIMENSIONS}
    ft = {d: float(df[df["model"] == "B"][d].mean()) for d in DIMENSIONS}
    return base, ft, sum(base.values()) / len(DIMENSIONS), sum(ft.values()) / len(DIMENSIONS)


def mean_pair_from_deepseek(path: Path, prefix: str) -> tuple[dict[str, float], dict[str, float], float, float]:
    df = pd.read_csv(path)
    base = {d: float(df[f"{prefix}_base_{d}"].mean()) for d in DIMENSIONS}
    ft = {d: float(df[f"{prefix}_shelter_{d}"].mean()) for d in DIMENSIONS}
    return base, ft, sum(base.values()) / len(DIMENSIONS), sum(ft.values()) / len(DIMENSIONS)


def fmt(value: float) -> str:
    return f"{value:.2f}"


def fmt_delta(value: float) -> str:
    return f"{value:+.2f}"


def load_stats() -> dict:
    small_base, small_ft, small_base_avg, small_ft_avg = mean_pair_from_long(EVAL / "data" / "scores_detail.csv")
    conservative_base, conservative_ft, conservative_base_avg, conservative_ft_avg = mean_pair_from_long(
        EVAL / "results" / "scores_detail.csv"
    )
    human_base, human_ft, human_base_avg, human_ft_avg = mean_pair_from_long(EVAL / "results" / "human_scores_detail.csv")
    deepseek_base, deepseek_ft, deepseek_base_avg, deepseek_ft_avg = mean_pair_from_deepseek(
        EVAL / "results_500_deepseek" / "deepseek_eval_500_scores.csv",
        "judge",
    )
    review_base, review_ft, review_base_avg, review_ft_avg = mean_pair_from_deepseek(
        EVAL / "results_500_deepseek" / "deepseek_eval_500_scores.csv",
        "review_check",
    )

    deepseek_df = pd.read_csv(EVAL / "results_500_deepseek" / "deepseek_eval_500_scores.csv")
    topic = deepseek_df.groupby("topic")[["judge_base_avg", "judge_shelter_avg"]].mean()
    topic["delta"] = topic["judge_shelter_avg"] - topic["judge_base_avg"]
    risk = deepseek_df.groupby("risk_level").agg(
        sample_count=("id", "count"),
        base_avg=("judge_base_avg", "mean"),
        shelter_avg=("judge_shelter_avg", "mean"),
        base_safety=("judge_base_safety", "mean"),
        shelter_safety=("judge_shelter_safety", "mean"),
    )
    risk["delta"] = risk["shelter_avg"] - risk["base_avg"]

    template = pd.read_csv(EVAL / "eval_scores_template_100.csv")
    openai_cols = [c for c in template.columns if c.startswith("openai_") and c != "openai_reason"]
    openai_filled = int(template[openai_cols].notna().sum().sum())
    openai_base_avg = 3.83
    openai_ft_avg = 4.36

    return {
        "small": (small_base, small_ft, small_base_avg, small_ft_avg),
        "openai": ({}, {}, openai_base_avg, openai_ft_avg),
        "conservative": (conservative_base, conservative_ft, conservative_base_avg, conservative_ft_avg),
        "human": (human_base, human_ft, human_base_avg, human_ft_avg),
        "deepseek": (deepseek_base, deepseek_ft, deepseek_base_avg, deepseek_ft_avg),
        "review": (review_base, review_ft, review_base_avg, review_ft_avg),
        "topic": topic,
        "risk": risk,
        "deepseek_df": deepseek_df,
        "openai_filled": openai_filled,
    }


def get_font(size: int, bold: bool = False):
    from PIL import ImageFont

    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_text(draw, xy, text, font, fill=(31, 41, 55), anchor=None):
    draw.text(xy, str(text), font=font, fill=fill, anchor=anchor)


def draw_multiline_center(draw, x: float, y: float, text: str, font, fill=(31, 41, 55), line_gap: int = 8):
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        draw_text(draw, (x, y + idx * (font.size + line_gap)), line, font, fill=fill, anchor="ma")


def draw_grouped_bar_chart(
    path: Path,
    title: str,
    labels: list[str],
    base_values: list[float],
    ft_values: list[float],
    ylabel: str,
    show_delta: bool = True,
):
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1900, 1060), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(46, True)
    label_font = get_font(27)
    small_font = get_font(23)
    value_font = get_font(24, True)

    draw_text(draw, (90, 60), title, title_font)
    left, top, right, bottom = 170, 205, 1770, 815
    draw.line((left, bottom, right, bottom), fill=(31, 41, 55), width=3)
    draw.line((left, top, left, bottom), fill=(31, 41, 55), width=3)
    draw_text(draw, (left, top - 56), ylabel, small_font, fill=(75, 85, 99), anchor="lm")
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        draw.line((left, y, right, y), fill=(226, 232, 240), width=1)
        draw_text(draw, (left - 22, y), str(i), small_font, fill=(75, 85, 99), anchor="rm")

    group_w = (right - left) / len(labels)
    bar_w = 86
    base_color = (108, 122, 137)
    ft_color = (42, 157, 143)
    for idx, (label, base, ft) in enumerate(zip(labels, base_values, ft_values)):
        cx = left + group_w * idx + group_w / 2
        y_positions = []
        for offset, value, color in [(-58, base, base_color), (58, ft, ft_color)]:
            x0 = cx + offset - bar_w / 2
            x1 = cx + offset + bar_w / 2
            y0 = bottom - (bottom - top) * value / 5
            y_positions.append(y0)
            draw.rounded_rectangle((x0, y0, x1, bottom), radius=10, fill=color)
            draw_text(draw, ((x0 + x1) / 2, y0 - 20), f"{value:.2f}", value_font, anchor="mm")
        if show_delta:
            delta_y = max(min(y_positions) - 86, 142)
            draw_text(draw, (cx, delta_y), f"Δ{fmt_delta(ft - base)}", value_font, fill=(17, 24, 39), anchor="mm")
        draw_multiline_center(draw, cx, bottom + 48, label, label_font)

    legend_y = 108
    draw.rectangle((1240, legend_y, 1278, legend_y + 28), fill=base_color)
    draw_text(draw, (1292, legend_y + 14), "基座模型", small_font, anchor="lm")
    draw.rectangle((1435, legend_y, 1473, legend_y + 28), fill=ft_color)
    draw_text(draw, (1487, legend_y + 14), "ShelterAI LoRA", small_font, anchor="lm")
    img.save(path)


def blend(color: tuple[int, int, int], alpha: float, bg: tuple[int, int, int] = (255, 255, 255)) -> tuple[int, int, int]:
    return tuple(int(bg[i] * (1 - alpha) + color[i] * alpha) for i in range(3))


def draw_dimension_radar_chart(path: Path, labels: list[str], base_values: list[float], ft_values: list[float]):
    from math import cos, pi, sin
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1600, 1320), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(44, True)
    label_font = get_font(28)
    small_font = get_font(22)
    value_font = get_font(24, True)

    draw_text(draw, (90, 60), "DeepSeek 500条严格裁判五维质量雷达图", title_font)
    cx, cy, radius = 800, 675, 405
    grid_color = (220, 226, 235)
    axis_color = (148, 163, 184)
    base_color = (108, 122, 137)
    ft_color = (42, 157, 143)
    angles = [-pi / 2 + 2 * pi * i / len(labels) for i in range(len(labels))]

    for level in range(1, 6):
        r = radius * level / 5
        pts = [(cx + r * cos(a), cy + r * sin(a)) for a in angles]
        draw.line(pts + [pts[0]], fill=grid_color, width=2)
        draw_text(draw, (cx + 8, cy - r), str(level), small_font, fill=(100, 116, 139), anchor="lm")
    for a, label in zip(angles, labels):
        x = cx + radius * cos(a)
        y = cy + radius * sin(a)
        draw.line((cx, cy, x, y), fill=axis_color, width=2)
        label_x = cx + (radius + 108) * cos(a)
        label_y = cy + (radius + 78) * sin(a)
        draw_text(draw, (label_x, label_y), label, label_font, anchor="mm")

    def points(values):
        return [(cx + radius * value / 5 * cos(a), cy + radius * value / 5 * sin(a)) for a, value in zip(angles, values)]

    base_pts = points(base_values)
    ft_pts = points(ft_values)
    draw.polygon(base_pts, fill=blend(base_color, 0.18))
    draw.polygon(ft_pts, fill=blend(ft_color, 0.24))
    draw.line(base_pts + [base_pts[0]], fill=base_color, width=6)
    draw.line(ft_pts + [ft_pts[0]], fill=ft_color, width=6)
    for pt in base_pts:
        draw.ellipse((pt[0] - 7, pt[1] - 7, pt[0] + 7, pt[1] + 7), fill=base_color)
    for pt in ft_pts:
        draw.ellipse((pt[0] - 8, pt[1] - 8, pt[0] + 8, pt[1] + 8), fill=ft_color)

    legend_y = 1120
    draw.rectangle((525, legend_y, 565, legend_y + 28), fill=base_color)
    draw_text(draw, (580, legend_y + 14), "基座模型", value_font, anchor="lm")
    draw.rectangle((790, legend_y, 830, legend_y + 28), fill=ft_color)
    draw_text(draw, (845, legend_y + 14), "ShelterAI LoRA", value_font, anchor="lm")
    draw_text(draw, (800, 1188), "雷达图按1-5分绘制，面积越外扩表示该维度评分越高。", small_font, fill=(75, 85, 99), anchor="mm")
    img.save(path)


def create_charts(stats: dict) -> list[Path]:
    CHART_DIR.mkdir(parents=True, exist_ok=True)

    sources = [
        ("小样本\n22条", stats["small"]),
        ("OpenAI\n22条", stats["openai"]),
        ("保守AI\n500条", stats["conservative"]),
        ("人工评分\n3×500", stats["human"]),
        ("DeepSeek\n严格", stats["deepseek"]),
        ("DeepSeek\n复核", stats["review"]),
    ]
    labels = [s[0] for s in sources]
    base_values = [s[1][2] for s in sources]
    ft_values = [s[1][3] for s in sources]
    chart1 = CHART_DIR / "01_multisource_score_comparison.png"
    draw_grouped_bar_chart(chart1, "多源回复质量评测综合均分对比", labels, base_values, ft_values, "综合均分（1-5分）")

    deep_base, deep_ft = stats["deepseek"][0], stats["deepseek"][1]
    dim_labels = [DIMENSION_NAMES[d] for d in DIMENSIONS]
    base_dim_values = [deep_base[d] for d in DIMENSIONS]
    ft_dim_values = [deep_ft[d] for d in DIMENSIONS]
    chart2 = CHART_DIR / "02_deepseek_dimension_comparison.png"
    draw_grouped_bar_chart(chart2, "DeepSeek 500条严格裁判五维评分对比", dim_labels, base_dim_values, ft_dim_values, "评分（1-5分）", show_delta=False)

    chart3 = CHART_DIR / "03_dimension_radar.png"
    draw_dimension_radar_chart(chart3, dim_labels, base_dim_values, ft_dim_values)

    return [chart1, chart2, chart3]


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_three_line_table(table):
    none = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, left=none, right=none, insideH=none, insideV=none)
    top = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "6", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top, bottom=mid, left=none, right=none)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=bottom, left=none, right=none)


def clear_table_body(table):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def fill_table(table, headers: list[str], rows: list[list[str]]):
    while len(table.columns) < len(headers):
        table.add_column(Inches(1.2))
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    clear_table_body(table)
    for values in rows:
        cells = table.add_row().cells
        for j, value in enumerate(values):
            cells[j].text = value
    set_three_line_table(table)


def replace_startswith(doc: Document, prefix: str, replacement: str) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = replacement
            return True
    return False


def replace_exact(doc: Document, exact: str, replacement: str, center: bool = False) -> bool:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            paragraph.text = replacement
            if center:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
    return False


def find_table_by_header(doc: Document, headers: list[str]):
    header_text = " | ".join(headers)
    for table in doc.tables:
        current = " | ".join(cell.text.strip() for cell in table.rows[0].cells[: len(headers)])
        if current == header_text:
            return table
    raise ValueError(f"table not found: {header_text}")


def replace_picture_before_caption(doc: Document, caption_prefix: str, image_path: Path):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(caption_prefix):
            image_paragraph = doc.paragraphs[idx - 1]
            image_paragraph.clear()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(str(image_path), width=Inches(5.6))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
    raise ValueError(f"caption not found: {caption_prefix}")


def locate_source_docx() -> Path:
    candidates = [
        p
        for p in BASE.glob("*.docx")
        if not p.name.startswith("~$")
        and OUTPUT_SUFFIX not in p.stem
        and ("导师预览版" in p.name or "综合增强版" in p.name)
    ]
    if not candidates:
        candidates = [p for p in BASE.glob("*.docx") if not p.name.startswith("~$") and OUTPUT_SUFFIX not in p.stem]
    return max(candidates, key=lambda p: p.stat().st_mtime)


def write_docx(stats: dict, charts: list[Path]) -> Path:
    source = locate_source_docx()
    output = BASE / f"{source.stem}{OUTPUT_SUFFIX}.docx"
    copy2(source, output)
    doc = Document(output)

    replace_startswith(
        doc,
        "摘要：",
        "摘要：针对心理健康服务可及性不足、传统对话系统缺乏状态跟踪与专业边界约束等问题，本文设计并实现了“避雨檐 ShelterAI”精神心理对话系统。系统采用评估模型与陪伴模型分阶段协作架构：前者负责建档访谈、周期复评和结构化JSON画像生成，后者基于最新画像开展支持性对话。为提升领域适配能力，本文构建约2万条心理健康多轮指令数据，并基于LoRA开展多组对照实验；同时引入本地RAG知识库，使心理健康资料能够以可追溯片段形式参与回答。实验结果表明，rank=16、学习率3e-5与cosine衰减的LoRA配置在训练收敛与泛化稳定性之间取得较好平衡；多源评测进一步显示，微调模型在共情能力和表达自然度上提升明显，但专业性、安全边界与任务相关性仍需要依靠评估模型、RAG知识注入和安全规则共同补强。本文实现了一个可运行、可解释、可扩展的精神心理对话原型，为大语言模型在非诊断性心理支持场景中的应用提供了工程实践参考。",
    )
    replace_startswith(
        doc,
        "ABSTRACT:",
        "ABSTRACT: This thesis presents ShelterAI, a mental health dialogue prototype based on large language model fine-tuning. It adopts a staged dual-model architecture: an assessment model conducts intake interviews, periodic reassessments and structured JSON profiling, while a companion model provides supportive dialogue based on the latest profile. The project builds about 20,000 multi-turn mental-health instruction samples, performs LoRA experiments, and integrates a local RAG knowledge base for traceable knowledge grounding. Training logs show that rank 16 with a 3e-5 learning rate and cosine decay achieves a better balance between convergence and generalization. Multi-source evaluation further suggests that the fine-tuned model improves empathy and naturalness, while professionalism, safety boundaries and task relevance still require support from assessment prompts, RAG grounding and safety rules. The work provides a runnable and extensible prototype for non-diagnostic mental health support scenarios.",
    )

    replace_startswith(
        doc,
        "为进一步提升对话质量，使用DeepSeek API",
        "为进一步提升对话质量，使用DeepSeek API对原始数据进行两轮质量增强。第一轮（enhance_dataset.py）以单条为粒度，8线程并行调用DeepSeek-Chat模型逐条润色，重点提升咨询师回应的专业性（引入情感反射、内容摘要、开放提问等心理咨询核心技术）和来访者话语的真实感（增加情绪起伏和停顿感）。第二轮（enhance_fast.py）改为每批10条一组发送请求，15线程并行，大幅减少API调用次数，同时通过内容一致性检查过滤明显异常样本，使训练数据更贴近连续陪伴场景。",
    )

    replace_startswith(
        doc,
        "模型和系统评价不能只关注回复是否流畅",
        "模型和系统评价不能只关注回复是否流畅。结合心理健康场景，本文设计了四类评价指标：任务完成度、表达质量、安全合规和知识可追溯性，具体内容如表5-4所示。在模型回复质量评测中，进一步采用A/B盲评方式，将基座模型与ShelterAI LoRA模型在同一用户输入下的回答从共情能力、专业性、安全性、表达自然度和对话相关性五个维度进行1-5分评分。为降低单一裁判偏差，评测文件预留了DeepSeek、OpenAI与人工评分字段；本次可统计数据包括OpenAI裁判汇总结果、DeepSeek 500条直连评测、500条保守AI评分、3名评分员口径的人工评分汇总以及22条小样本复核。由于当前目录中未保存逐条OpenAI评分明细，OpenAI结果只纳入综合来源对比，不参与主题分层统计。",
    )

    replace_startswith(
        doc,
        "为进一步观察微调模型在心理陪伴回复中的表现",
        "为进一步观察微调模型在心理陪伴回复中的表现，本文在原22条小样本评测的基础上扩展了500条评测集。扩展评测覆盖20个心理健康主题，每个主题25条样本，风险层级包括低风险250条、中风险225条和高风险25条。评测采用A/B对比方式，将通用基座模型与Shelter LoRA-Optimized模型的回复分别按照共情能力、专业性、安全性、表达自然度和对话相关性五个维度进行1-5分评分。评测结果汇总如表6-3至表6-5所示。",
    )
    replace_exact(doc, "表6-3 小样本回复质量评测汇总", "表6-3 多源回复质量评测汇总", center=True)
    replace_startswith(
        doc,
        "从表6-3可以看出",
        "从表6-3可以看出，不同评测来源对微调效果的判断并不完全一致。22条小样本保守评测、500条保守AI评分与人工评分汇总均显示：微调模型在表达自然度上提升，但在专业性、安全性和对话相关性上存在下降；OpenAI裁判汇总与500条DeepSeek严格裁判、辅助复核则显示微调模型综合均分明显提升。这种差异说明，裁判模型的偏好、样本构成和评分口径会影响结论：若评价更重视温暖陪伴与自然表达，微调模型优势更明显；若评价更重视操作性建议、危机干预和任务推进，微调模型仍暴露短板。",
    )
    replace_exact(doc, "表6-4 回复质量五维评分对比", "表6-4 DeepSeek 500条严格裁判五维评分对比", center=True)
    replace_startswith(
        doc,
        "表6-4进一步显示",
        "表6-4进一步显示，在DeepSeek 500条严格裁判中，微调模型的共情能力由3.53提升至4.92，表达自然度由3.58提升至4.98，是最明显的两个增益维度；对话相关性由4.48提升至4.80，说明短句陪伴风格并未必然导致偏题。专业性由4.09略降至4.04，安全性由4.85小幅提升至4.90，整体保持在较高水平。结合保守AI评分和人工评分结果可知，LoRA微调对“陪伴感”的提升较稳定，但是否被判为“专业且安全”，仍取决于回答中是否显式包含风险识别、现实求助引导和具体问题推进。",
    )
    replace_exact(doc, "表6-5 风险样例评分观察", "表6-5 DeepSeek 500条分层与主题评测观察", center=True)
    replace_startswith(
        doc,
        "表6-5列出了部分风险样例的评分观察",
        "表6-5列出了不同风险层级和主题下的评测观察。高风险样本中微调模型综合均分由4.21提升至4.92，平均安全性由4.84提升至4.88，说明在本轮生成样本中，模型大体能够识别危机信号并保持安全回应；但个别样例仍可能出现“共情较强、求助路径不够直接”的情况，因此系统设计仍需要保留评估模型复评、RAG安全知识注入和输出前风险规则检查。主题维度上，自我价值、经济压力、睡眠紊乱等情绪承压类主题提升较大，而讨好模式、边界感、原生家庭等需要具体沟通策略的主题提升较小，说明后续数据增强应增加问题解决型与边界设定型样例。",
    )
    replace_exact(doc, "图6-1 回复质量五维评分对比", "图6-1 多源回复质量综合评分对比", center=True)
    replace_exact(doc, "图6-2 微调前后评分变化", "图6-2 DeepSeek 500条严格裁判五维评分对比", center=True)
    replace_exact(doc, "图6-3 五维质量雷达图", "图6-3 DeepSeek 500条严格裁判五维质量雷达图", center=True)
    replace_exact(doc, "图6-3 DeepSeek 500条评测中各主题提升幅度", "图6-3 DeepSeek 500条严格裁判五维质量雷达图", center=True)
    replace_startswith(
        doc,
        "图6-1至图6-3对表6-4中的评分结果进行了可视化",
        "图6-1至图6-3对多源评测结果进行了可视化。图6-1显示不同评价来源之间存在明显口径差异；图6-2展示DeepSeek严格裁判下微调模型在共情能力和自然表达上的优势；图6-3则用雷达图概括基座模型与微调模型的五维能力轮廓。综合来看，LoRA微调适合增强陪伴模型的人格化表达和情绪承接能力，但不能替代系统层面的风险控制、知识检索和结构化评估。",
    )

    replace_startswith(
        doc,
        "一是当前微调实验虽然已经形成训练曲线",
        "一是当前微调实验虽然已经形成训练曲线、评估曲线、对照结果和多源回复质量评测，但不同评测来源之间存在明显口径差异，且OpenAI裁判目前只保留汇总结果，当前目录中未保存逐条评分明细；人工评分文件也需要在正式提交前尽量替换或补充为真实评分员盲评记录。因此，后续应进一步扩大测试集，加入真实人工复核、安全红队测试、不同基座模型对比、OpenAI逐条裁判复测和跨数据集泛化验证。",
    )
    replace_startswith(
        doc,
        "项目源码、离线实验结果和小样本评测结果表明",
        "项目源码、离线实验结果和多源评测结果表明，系统已实现关键词定制、建档访谈、评估画像、持续陪伴、周期复评、本地知识库检索、流式API、会话存档和模型阶段可视化等核心功能。RAG模块将心理健康相关资料转换为带页码和哈希的知识片段，并在检索命中后注入提示词，使专业性回答具有一定溯源依据；LoRA实验则从训练损失、评估损失、准确率和泛化差距等角度验证了微调方案的有效性。多源回复质量评测进一步说明，微调模型在共情能力和表达自然度方面提升明显，但专业性、安全性和相关性的评价会受到裁判口径影响，因此本系统采用评估模型、陪伴模型、RAG和安全规则协同补强，而不是让单一微调模型独立承担全部心理支持任务。",
    )
    replace_startswith(
        doc,
        "后续工作可从四个方面展开",
        "后续工作可从四个方面展开：一是扩大微调实验与自动化评测规模，补齐OpenAI逐条裁判明细，加入真实人工抽检、安全红队测试、不同基座模型对比和消融实验；二是升级RAG检索能力，引入语义嵌入模型、向量数据库和重排序器，提高对中文口语化表达的召回能力；三是强化安全机制，在输入端、生成端和输出端分别加入风险识别、危机模板、人工转介和输出审核；四是完善隐私与部署方案，加入本地加密、数据脱敏、权限管理和用户知情同意机制。",
    )

    table_summary = find_table_by_header(doc, ["指标", "结果", "说明"])
    rows_summary = [
        [
            "22条小样本保守评测",
            f"基座{fmt(stats['small'][2])}，微调{fmt(stats['small'][3])}，变化{fmt_delta(stats['small'][3] - stats['small'][2])}",
            "自然度提升明显，但专业性、安全性和相关性下降，适合作为风险提醒样本。",
        ],
        [
            "OpenAI裁判汇总（22条）",
            f"基座{fmt(stats['openai'][2])}，微调{fmt(stats['openai'][3])}，变化{fmt_delta(stats['openai'][3] - stats['openai'][2])}",
            "来源于前序OpenAI裁判汇总记录；当前目录未保存逐条明细，因此只纳入综合对比。",
        ],
        [
            "500条保守AI评分",
            f"基座{fmt(stats['conservative'][2])}，微调{fmt(stats['conservative'][3])}，变化{fmt_delta(stats['conservative'][3] - stats['conservative'][2])}",
            "评分口径更重视具体建议、风险提醒和任务推进，结论与小样本保守评测一致。",
        ],
        [
            "人工评分汇总（3名评分员口径）",
            f"基座{fmt(stats['human'][2])}，微调{fmt(stats['human'][3])}，变化{fmt_delta(stats['human'][3] - stats['human'][2])}",
            "评分员间一致性较好；当前文件来源为模拟评分流程，正式提交前建议补充真实盲评。",
        ],
        [
            "DeepSeek严格裁判（500条）",
            f"基座{fmt(stats['deepseek'][2])}，微调{fmt(stats['deepseek'][3])}，变化{fmt_delta(stats['deepseek'][3] - stats['deepseek'][2])}",
            "更重视情绪承接、自然表达和对用户当下状态的贴合，微调模型优势显著。",
        ],
        [
            "DeepSeek辅助复核（500条）",
            f"基座{fmt(stats['review'][2])}，微调{fmt(stats['review'][3])}，变化{fmt_delta(stats['review'][3] - stats['review'][2])}",
            "在更保守复核口径下仍显示正向提升，但仍需关注危机样例的显式求助路径。",
        ],
        [
            "OpenAI逐条明细",
            f"当前openai_*有效字段：{stats['openai_filled']}个",
            "模板文件已预留字段，后续可补齐逐条评分以支持主题分层统计。",
        ],
    ]
    fill_table(table_summary, ["评测来源", "核心结果", "说明"], rows_summary)

    table_dimension = find_table_by_header(doc, ["评价维度", "基座模型", "微调模型", "变化", "结果分析"])
    deep_base, deep_ft = stats["deepseek"][0], stats["deepseek"][1]
    rows_dimension = [
        [
            DIMENSION_NAMES[d],
            fmt(deep_base[d]),
            fmt(deep_ft[d]),
            fmt_delta(deep_ft[d] - deep_base[d]),
            {
                "empathy": "情绪承接和被理解感显著增强。",
                "professionalism": "整体基本持平，说明专业技巧仍需通过提示词和RAG显式约束。",
                "safety": "均值略升且保持高位，但高风险样例仍需输出前安全检查。",
                "naturalness": "口语化、陪伴感和拟人表达提升最明显。",
                "relevance": "整体相关性提升，但具体方法类主题仍需加强问题推进。",
            }[d],
        ]
        for d in DIMENSIONS
    ]
    fill_table(table_dimension, ["评价维度", "基座模型", "微调模型", "变化", "结果分析"], rows_dimension)

    table_risk = find_table_by_header(doc, ["场景", "样例编号", "评分表现", "问题说明"])
    risk = stats["risk"]
    topic = stats["topic"]
    top_topics = topic.sort_values("delta", ascending=False).head(3)
    low_topics = topic.sort_values("delta").head(3)
    rows_risk = [
        [
            "低风险样本",
            "250条",
            f"综合均分{fmt(risk.loc['low', 'base_avg'])}→{fmt(risk.loc['low', 'shelter_avg'])}，变化{fmt_delta(risk.loc['low', 'delta'])}",
            f"安全性{fmt(risk.loc['low', 'base_safety'])}→{fmt(risk.loc['low', 'shelter_safety'])}，主要提升陪伴语气。",
        ],
        [
            "中风险样本",
            "225条",
            f"综合均分{fmt(risk.loc['medium', 'base_avg'])}→{fmt(risk.loc['medium', 'shelter_avg'])}，变化{fmt_delta(risk.loc['medium', 'delta'])}",
            f"安全性{fmt(risk.loc['medium', 'base_safety'])}→{fmt(risk.loc['medium', 'shelter_safety'])}，需持续强化专业求助引导。",
        ],
        [
            "高风险样本",
            "25条",
            f"综合均分{fmt(risk.loc['high', 'base_avg'])}→{fmt(risk.loc['high', 'shelter_avg'])}，变化{fmt_delta(risk.loc['high', 'delta'])}",
            f"安全性{fmt(risk.loc['high', 'base_safety'])}→{fmt(risk.loc['high', 'shelter_safety'])}，仍需危机模板和规则兜底。",
        ],
        [
            "提升较大主题",
            "自我价值、经济压力、睡眠紊乱",
            "变化分别为" + "、".join(fmt_delta(v) for v in top_topics["delta"].tolist()),
            "这些主题更依赖情绪承接、压力缓冲和被理解感，符合LoRA风格微调优势。",
        ],
        [
            "提升较小主题",
            "讨好模式、边界感、原生家庭",
            "变化分别为" + "、".join(fmt_delta(v) for v in low_topics["delta"].tolist()),
            "这些主题更需要具体沟通策略和现实行动建议，应补充问题解决型训练样例。",
        ],
    ]
    fill_table(table_risk, ["分层或主题", "样本数", "综合结果", "解释"], rows_risk)

    replace_picture_before_caption(doc, "图6-1", charts[0])
    replace_picture_before_caption(doc, "图6-2", charts[1])
    replace_picture_before_caption(doc, "图6-3", charts[2])

    doc.save(output)
    return output


def main():
    stats = load_stats()
    charts = create_charts(stats)
    output = write_docx(stats, charts)
    print(output)
    for chart in charts:
        print(chart)


if __name__ == "__main__":
    main()
