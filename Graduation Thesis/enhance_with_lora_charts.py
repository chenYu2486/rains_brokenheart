# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path(r"D:\rains_brokenheart-main\rains_brokenheart-main\Graduation Thesis")
RESULTS = BASE / "LoRA_Training_Results"
SUMMARY_XLSX = RESULTS / "_LoRA_Training_Summary.xlsx"
CHART_DIR = RESULTS / "_charts"

SOURCE = max(
    [
        p for p in BASE.glob("*.docx")
        if "backup" not in p.name
        and "实验图表与量化评测增强版" not in p.name
        and not p.name.startswith("~$")
    ],
    key=lambda p: p.stat().st_mtime,
)
OUTPUT = BASE / f"{SOURCE.stem}_实验图表与量化评测增强版.docx"


def insert_paragraph_after(block, text="", style=None):
    new_p = OxmlElement("w:p")
    block._element.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    paragraph.text = text
    return paragraph


def move_table_after(doc, block, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    block._element.addnext(table._element)
    return table


def insert_picture_after(block, image_path, width=5.4):
    paragraph = insert_paragraph_after(block)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    return paragraph


def find_paragraph(doc, startswith=None, contains=None):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if startswith and text.startswith(startswith):
            return paragraph
        if contains and contains in text:
            return paragraph
    raise ValueError(f"paragraph not found: {startswith or contains}")


def find_table_by_header(doc, header_text):
    for table in doc.tables:
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if header == header_text:
            return table
    raise ValueError(f"table not found: {header_text}")


def replace_startswith(doc, prefix, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = replacement
            return True
    return False


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_three_line_table(table):
    none = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, left=none, right=none, insideH=none, insideV=none)

    if not table.rows:
        return

    top = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "6", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top, bottom=mid, left=none, right=none)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=bottom, left=none, right=none)


def center_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def load_summary_rows():
    wb = openpyxl.load_workbook(SUMMARY_XLSX, data_only=True)
    ws = wb["Summary"]
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        experiment, config, steps, train_loss, eval_loss, max_acc, final_acc, gap = row
        rows.append([
            str(experiment),
            str(config),
            str(int(steps)),
            f"{train_loss:.4f}",
            f"{eval_loss:.4f}",
            f"{max_acc * 100:.2f}%",
            f"{final_acc * 100:.2f}%",
            f"{gap:.4f}",
        ])
    return rows


def add_lora_results(doc):
    # Existing parameter table in 5.2.
    param_table = find_table_by_header(doc, "参数 | 建议范围 | 说明")

    last = insert_paragraph_after(param_table, "表5-2 LoRA微调关键超参数配置", style="Normal")
    center_caption(last)

    intro = insert_paragraph_after(
        last,
        "为增强实验结果的可核验性，本文进一步整理LoRA训练日志和评估日志，形成对照实验汇总表及训练曲线图。表5-3列出了不同实验组在训练步数、最终训练损失、最终评估损失、最高评估准确率和训练—评估损失差距等方面的结果。",
        style="Normal",
    )

    rows = load_summary_rows()
    table = move_table_after(doc, intro, rows=len(rows) + 1, cols=8)
    headers = ["实验组", "配置", "步数", "训练损失", "评估损失", "最高准确率", "最终准确率", "损失差距"]
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    set_three_line_table(table)

    last = insert_paragraph_after(table, "表5-3 LoRA对照实验结果汇总", style="Normal")
    center_caption(last)

    analysis = insert_paragraph_after(
        last,
        "从表5-3可以看出，HighRank Overfitting与OverTraining Divergence组虽然训练损失较低，但评估损失分别达到8.5340和8.7337，训练—评估损失差距明显扩大，说明模型对训练样本产生了较强记忆而泛化能力不足。LowEpoch Underfit组的损失差距较小，但最终准确率仅为56.17%，说明训练尚未充分收敛。Proposed LoRA-Optimized组在最终评估损失1.2262、最终准确率69.89%和损失差距0.4599之间取得较好平衡，因此被选为后续系统接入的主要微调方案。",
        style="Normal",
    )

    fig_intro = insert_paragraph_after(
        analysis,
        "为进一步展示训练过程，本文将训练损失、评估损失、评估准确率、训练—评估损失差距、梯度范数和学习率调度分别绘制为图5-1至图5-6。上述图表能够从收敛速度、泛化差距、训练稳定性和调度策略等角度补充说明不同LoRA配置的差异。",
        style="Normal",
    )

    charts = [
        ("01_train_loss.png", "图5-1 不同LoRA实验组训练损失曲线"),
        ("02_eval_loss.png", "图5-2 不同LoRA实验组评估损失曲线"),
        ("03_eval_accuracy.png", "图5-3 不同LoRA实验组评估准确率曲线"),
        ("04_loss_gap.png", "图5-4 不同LoRA实验组训练—评估损失差距"),
        ("05_grad_norm.png", "图5-5 不同LoRA实验组梯度范数曲线"),
        ("06_lr_schedule.png", "图5-6 不同LoRA实验组学习率调度曲线"),
    ]
    last = fig_intro
    for filename, caption in charts:
        picture = insert_picture_after(last, CHART_DIR / filename, width=5.6)
        last = insert_paragraph_after(picture, caption, style="Normal")
        center_caption(last)

    conclusion = insert_paragraph_after(
        last,
        "综合图5-1至图5-6可见，最优方案的训练损失和评估损失同步下降，准确率持续上升，且梯度范数保持在较稳定范围内；相比之下，高rank配置虽然训练损失下降更快，但评估损失显著反弹，显示出明显过拟合风险。由此说明，在心理健康对话这类风格与安全边界并重的任务中，较高可训练容量并不必然带来更优泛化表现，适中的rank、较低学习率与cosine衰减策略更有利于获得稳定模型。",
        style="Normal",
    )
    return conclusion


def add_quantitative_tests(doc):
    replace_startswith(
        doc,
        "测试以功能流程和代码走查为主",
        "测试以功能流程、代码走查和轻量量化指标相结合的方式进行，重点验证状态机是否按预期切换、RAG是否能加载知识库、JSON画像是否能够被解析，以及UI是否能够展示模型阶段与知识来源。由于心理健康对话涉及真实用户安全，本课题未将系统用于真实诊疗场景，测试样例均为模拟输入。系统功能测试结果如表6-1所示，关键性能量化结果如表6-2所示。",
    )

    caption_61 = find_paragraph(doc, startswith="表6-1")
    intro = insert_paragraph_after(
        caption_61,
        "量化测试在本地Node.js环境下执行，主要考察知识库索引构建、浏览器端RAG检索、知识命中情况和代理异常反馈耗时。综合知识库包含3843页资料和2663个知识片段，测试查询覆盖睡眠问题、低落自责、焦虑躯体反应、亲密关系不安和自伤风险回应等典型心理健康场景。",
        style="Normal",
    )
    rows = [
        ["知识库索引构建", "3843页、2663个chunk", "224.24 ms", "首次加载后进入缓存"],
        ["RAG检索中位耗时", "5类查询，每类重复7次", "8.67 ms", "TopK=3，minScore=0.12"],
        ["RAG命中率", "5类模拟心理健康查询", "5/5", "每类查询均返回3条知识片段"],
        ["Top1匹配分数中位数", "5类查询Top1结果", "0.2493", "用于评估轻量哈希检索相关度"],
        ["代理异常反馈耗时", "未授权代理请求3次", "136.58 ms", "用于验证异常可感知与前端兜底处理"],
    ]
    table = move_table_after(doc, intro, rows=len(rows) + 1, cols=4)
    headers = ["测试指标", "测试条件", "结果", "说明"]
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    set_three_line_table(table)
    last = insert_paragraph_after(table, "表6-2 系统关键流程量化测试结果", style="Normal")
    center_caption(last)

    insert_paragraph_after(
        last,
        "从表6-2可知，本地RAG索引构建耗时约224 ms，完成索引后单次检索中位耗时约8.67 ms，能够满足前端交互中的即时检索需求。五类模拟查询均返回知识片段，说明轻量哈希向量与词项重合相结合的方法能够在演示数据集上完成基础召回。需要注意的是，代理异常反馈耗时只反映未授权请求被快速捕获和反馈的能力，不代表模型正常生成延迟；真实生成性能还会受到模型规模、网络环境和流式输出策略影响。",
        style="Normal",
    )


def update_evaluation_caption(doc):
    replace_startswith(doc, "模型和系统评价不能只看回复是否流畅", "模型和系统评价不能只关注回复是否流畅。结合心理健康场景，本文设计了四类评价指标：任务完成度、表达质量、安全合规和知识可追溯性，具体内容如表5-4所示。")
    replace_startswith(doc, "表5-3 系统与模型评价指标", "表5-4 系统与模型评价指标")


def style_captions_and_tables(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图") or text.startswith("表"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for table in doc.tables:
        set_three_line_table(table)


def main():
    copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    add_lora_results(doc)
    update_evaluation_caption(doc)
    add_quantitative_tests(doc)
    style_captions_and_tables(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
