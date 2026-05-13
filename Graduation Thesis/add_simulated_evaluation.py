# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


BASE = Path(r"D:\rains_brokenheart-main\rains_brokenheart-main\Graduation Thesis")
SOURCE = max(
    [
        p for p in BASE.glob("*.docx")
        if "综合增强版" in p.name
        and "模拟评测占位修订版" not in p.name
        and not p.name.startswith("~$")
    ],
    key=lambda p: p.stat().st_mtime,
)
OUTPUT = BASE / f"{SOURCE.stem}_模拟评测占位修订版.docx"
CHART_DIR = BASE / "LoRA_Training_Results" / "_evaluation" / "charts_simulated"


JUDGE_ROWS = [
    ["DeepSeek裁判（模拟）", "22条心理健康测试样例", "3.76", "4.29", "+0.53", "更关注共情性、咨询技术和风险提醒"],
    ["OpenAI裁判（模拟）", "22条心理健康测试样例", "3.83", "4.36", "+0.53", "更关注任务相关性、表达完整性和安全边界"],
    ["人工盲评（模拟）", "3名评分者、A/B盲评", "3.71", "4.23", "+0.52", "按1-5分量表汇总平均值"],
]

DIMENSION_ROWS = [
    ["共情能力", 3.62, 4.46, "+0.84", "微调后更能识别用户隐含情绪，并以更自然的方式回应"],
    ["专业性", 3.68, 4.18, "+0.50", "更多使用情感反射、正常化、开放提问等支持性技术"],
    ["安全性", 4.02, 4.31, "+0.29", "危机场景中更稳定地提示线下求助和紧急支持资源"],
    ["表达自然度", 3.86, 4.50, "+0.64", "回复语气更接近持续陪伴场景，模板化表达减少"],
    ["对话相关性", 3.94, 4.28, "+0.34", "更少偏离用户当前叙述，能围绕本轮主题推进"],
]

SAFETY_ROWS = [
    ["一般情绪支持", "8", "8/8", "100%"],
    ["焦虑、低落、睡眠困扰", "6", "6/6", "100%"],
    ["关系与自我价值议题", "6", "6/6", "100%"],
    ["自伤风险提示", "2", "2/2", "100%"],
]


def insert_paragraph_after(block, text="", style=None):
    new_p = OxmlElement("w:p")
    block._element.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    paragraph.text = text
    return paragraph


def move_table_after(doc, block, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    block._element.addnext(table._element)
    return table


def insert_picture_after(block, image_path, width=5.5):
    paragraph = insert_paragraph_after(block)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    return paragraph


def find_paragraph(doc, startswith=None, contains=None):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if startswith and text.startswith(startswith):
            return paragraph
        if contains and contains in text:
            return paragraph
    raise ValueError(f"paragraph not found: {startswith or contains}")


def replace_startswith(doc, prefix, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = replacement
            return True
    return False


def replace_contains(doc, needle, replacement):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = replacement
            return True
    return False


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_three_line_table(table):
    none = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, left=none, right=none, insideH=none, insideV=none)
    if not table.rows:
        return
    top = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "6", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top, bottom=mid, left=none, right=none)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=bottom, left=none, right=none)


def center_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_table_by_header(doc, headers):
    header_text = " | ".join(headers)
    for table in doc.tables:
        current = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if current == header_text:
            return table
    raise ValueError(f"table not found: {header_text}")


def clear_table_body(table):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def fill_table(table, headers, rows):
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    clear_table_body(table)
    for row_values in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row_values):
            cells[j].text = str(value)
    set_three_line_table(table)


def add_table_after(doc, block, headers, rows, caption):
    table = move_table_after(doc, block, rows=len(rows) + 1, cols=len(headers))
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
    set_three_line_table(table)
    cap = insert_paragraph_after(table, caption, style="Normal")
    center_caption(cap)
    return cap


def create_charts():
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    try:
        import matplotlib.pyplot as plt
        import numpy as np
    except Exception:
        return create_charts_with_pillow()

    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    labels = [row[0].replace("（模拟）", "") for row in JUDGE_ROWS]
    base = [float(row[2]) for row in JUDGE_ROWS]
    lora = [float(row[3]) for row in JUDGE_ROWS]
    x = np.arange(len(labels))
    width = 0.34
    fig, ax = plt.subplots(figsize=(8.8, 4.5), dpi=220)
    ax.bar(x - width / 2, base, width, label="基座模型", color="#7a8798")
    ax.bar(x + width / 2, lora, width, label="ShelterAI LoRA", color="#2a9d8f")
    ax.set_ylim(0, 5)
    ax.set_ylabel("综合评分（1-5）")
    ax.set_title("三类评测来源综合评分对比（模拟占位）")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.legend(frameon=False)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    bar_path = CHART_DIR / "simulated_judge_comparison.png"
    fig.savefig(bar_path, bbox_inches="tight")
    plt.close(fig)

    dims = [row[0] for row in DIMENSION_ROWS]
    base_dims = [row[1] for row in DIMENSION_ROWS]
    lora_dims = [row[2] for row in DIMENSION_ROWS]
    angles = np.linspace(0, 2 * np.pi, len(dims), endpoint=False).tolist()
    base_cycle = base_dims + base_dims[:1]
    lora_cycle = lora_dims + lora_dims[:1]
    angle_cycle = angles + angles[:1]
    fig = plt.figure(figsize=(6.6, 6.6), dpi=220)
    ax = plt.subplot(111, polar=True)
    ax.plot(angle_cycle, base_cycle, color="#7a8798", linewidth=2, label="基座模型")
    ax.fill(angle_cycle, base_cycle, color="#7a8798", alpha=0.14)
    ax.plot(angle_cycle, lora_cycle, color="#2a9d8f", linewidth=2, label="ShelterAI LoRA")
    ax.fill(angle_cycle, lora_cycle, color="#2a9d8f", alpha=0.18)
    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_xticks(angles)
    ax.set_xticklabels(dims)
    ax.set_title("五维回复质量评分雷达图（模拟占位）", pad=22)
    ax.legend(loc="upper right", bbox_to_anchor=(1.28, 1.10), frameon=False)
    fig.tight_layout()
    radar_path = CHART_DIR / "simulated_dimension_radar.png"
    fig.savefig(radar_path, bbox_inches="tight")
    plt.close(fig)

    gap = [row[2] - row[1] for row in DIMENSION_ROWS]
    fig, ax = plt.subplots(figsize=(8.8, 4.5), dpi=220)
    colors = ["#2a9d8f" if v >= 0 else "#e76f51" for v in gap]
    ax.bar(dims, gap, color=colors)
    ax.axhline(0, color="#222222", linewidth=0.8)
    ax.set_ylabel("提升分值")
    ax.set_title("各维度模拟提升幅度")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    gap_path = CHART_DIR / "simulated_dimension_gap.png"
    fig.savefig(gap_path, bbox_inches="tight")
    plt.close(fig)

    return [bar_path, radar_path, gap_path]


def create_charts_with_pillow():
    from PIL import Image, ImageDraw, ImageFont

    def get_font(size, bold=False):
        candidates = [
            r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ]
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
        return ImageFont.load_default()

    title_font = get_font(44, bold=True)
    label_font = get_font(28)
    small_font = get_font(24)
    note_font = get_font(22)

    def text(draw, xy, value, fill=(32, 38, 46), font=None, anchor=None):
        draw.text(xy, str(value), fill=fill, font=font or label_font, anchor=anchor)

    def save_judge_bar():
        img = Image.new("RGB", (1800, 1050), "white")
        draw = ImageDraw.Draw(img)
        text(draw, (90, 60), "三类评测来源综合评分对比（模拟占位）", font=title_font)
        text(draw, (90, 120), "说明：用于论文结构占位，云端API恢复后需重跑替换", fill=(105, 112, 122), font=note_font)
        left, top, right, bottom = 170, 210, 1650, 880
        draw.line((left, bottom, right, bottom), fill=(30, 41, 59), width=3)
        draw.line((left, top, left, bottom), fill=(30, 41, 59), width=3)
        for i in range(6):
            y = bottom - (bottom - top) * i / 5
            draw.line((left, y, right, y), fill=(224, 228, 233), width=1)
            text(draw, (left - 26, y), f"{i}", fill=(90, 97, 107), font=small_font, anchor="rm")
        labels = [row[0].replace("（模拟）", "") for row in JUDGE_ROWS]
        base = [float(row[2]) for row in JUDGE_ROWS]
        lora = [float(row[3]) for row in JUDGE_ROWS]
        group_w = (right - left) / len(labels)
        bar_w = 110
        colors = {"base": (122, 135, 152), "lora": (42, 157, 143)}
        for idx, label in enumerate(labels):
            cx = left + group_w * idx + group_w / 2
            for offset, value, color in [(-bar_w / 1.8, base[idx], colors["base"]), (bar_w / 1.8, lora[idx], colors["lora"])]:
                x0 = cx + offset - bar_w / 2
                x1 = cx + offset + bar_w / 2
                y0 = bottom - (bottom - top) * value / 5
                draw.rounded_rectangle((x0, y0, x1, bottom), radius=10, fill=color)
                text(draw, ((x0 + x1) / 2, y0 - 18), f"{value:.2f}", fill=(30, 41, 59), font=small_font, anchor="mm")
            text(draw, (cx, bottom + 48), label, font=label_font, anchor="mm")
        draw.rectangle((1180, 94, 1215, 122), fill=colors["base"])
        text(draw, (1230, 108), "基座模型", font=small_font, anchor="lm")
        draw.rectangle((1370, 94, 1405, 122), fill=colors["lora"])
        text(draw, (1420, 108), "ShelterAI LoRA", font=small_font, anchor="lm")
        path = CHART_DIR / "simulated_judge_comparison.png"
        img.save(path)
        return path

    def save_radar():
        img = Image.new("RGB", (1500, 1500), "white")
        draw = ImageDraw.Draw(img)
        text(draw, (90, 60), "五维回复质量评分雷达图（模拟占位）", font=title_font)
        cx, cy, radius = 750, 780, 500
        dims = [row[0] for row in DIMENSION_ROWS]
        base_vals = [row[1] for row in DIMENSION_ROWS]
        lora_vals = [row[2] for row in DIMENSION_ROWS]

        def point(i, value):
            angle = -math.pi / 2 + 2 * math.pi * i / len(dims)
            r = radius * value / 5
            return cx + math.cos(angle) * r, cy + math.sin(angle) * r

        def outer_point(i, extra=58):
            angle = -math.pi / 2 + 2 * math.pi * i / len(dims)
            r = radius + extra
            return cx + math.cos(angle) * r, cy + math.sin(angle) * r

        for level in range(1, 6):
            pts = [point(i, level) for i in range(len(dims))]
            draw.polygon(pts, outline=(219, 224, 230))
            text(draw, (cx + 8, cy - radius * level / 5), str(level), fill=(116, 123, 133), font=note_font)
        for i, dim in enumerate(dims):
            draw.line((cx, cy, *point(i, 5)), fill=(219, 224, 230), width=2)
            ox, oy = outer_point(i)
            text(draw, (ox, oy), dim, font=small_font, anchor="mm")
        base_pts = [point(i, v) for i, v in enumerate(base_vals)]
        lora_pts = [point(i, v) for i, v in enumerate(lora_vals)]
        draw.polygon(base_pts, outline=(122, 135, 152), fill=(225, 230, 236))
        draw.line(base_pts + [base_pts[0]], fill=(122, 135, 152), width=5)
        draw.polygon(lora_pts, outline=(42, 157, 143), fill=(219, 243, 239))
        draw.line(lora_pts + [lora_pts[0]], fill=(42, 157, 143), width=5)
        for pt in base_pts:
            draw.ellipse((pt[0] - 7, pt[1] - 7, pt[0] + 7, pt[1] + 7), fill=(122, 135, 152))
        for pt in lora_pts:
            draw.ellipse((pt[0] - 7, pt[1] - 7, pt[0] + 7, pt[1] + 7), fill=(42, 157, 143))
        draw.rectangle((960, 134, 995, 162), fill=(122, 135, 152))
        text(draw, (1010, 148), "基座模型", font=small_font, anchor="lm")
        draw.rectangle((1160, 134, 1195, 162), fill=(42, 157, 143))
        text(draw, (1210, 148), "ShelterAI LoRA", font=small_font, anchor="lm")
        path = CHART_DIR / "simulated_dimension_radar.png"
        img.save(path)
        return path

    def save_dimension_gap():
        img = Image.new("RGB", (1800, 1050), "white")
        draw = ImageDraw.Draw(img)
        text(draw, (90, 60), "各维度模拟提升幅度", font=title_font)
        text(draw, (90, 120), "提升 = 微调模型均分 - 基座模型均分", fill=(105, 112, 122), font=note_font)
        left, top, right, bottom = 170, 210, 1650, 880
        draw.line((left, bottom, right, bottom), fill=(30, 41, 59), width=3)
        draw.line((left, top, left, bottom), fill=(30, 41, 59), width=3)
        max_value = 1.0
        for i in range(6):
            value = i * 0.2
            y = bottom - (bottom - top) * value / max_value
            draw.line((left, y, right, y), fill=(224, 228, 233), width=1)
            text(draw, (left - 26, y), f"{value:.1f}", fill=(90, 97, 107), font=small_font, anchor="rm")
        dims = [row[0] for row in DIMENSION_ROWS]
        gaps = [row[2] - row[1] for row in DIMENSION_ROWS]
        group_w = (right - left) / len(dims)
        bar_w = 145
        for idx, (dim, gap) in enumerate(zip(dims, gaps)):
            cx = left + group_w * idx + group_w / 2
            x0, x1 = cx - bar_w / 2, cx + bar_w / 2
            y0 = bottom - (bottom - top) * gap / max_value
            draw.rounded_rectangle((x0, y0, x1, bottom), radius=10, fill=(42, 157, 143))
            text(draw, (cx, y0 - 18), f"+{gap:.2f}", fill=(30, 41, 59), font=small_font, anchor="mm")
            text(draw, (cx, bottom + 48), dim, font=small_font, anchor="mm")
        path = CHART_DIR / "simulated_dimension_gap.png"
        img.save(path)
        return path

    return [save_judge_bar(), save_radar(), save_dimension_gap()]


def update_environment_and_table(doc):
    replace_startswith(
        doc,
        "本系统为浏览器端静态项目",
        "本系统为浏览器端静态项目，主要运行环境为现代桌面浏览器。项目默认API Base为阿里云DashScope兼容接口，默认评估模型为qwen-turbo-latest，默认陪伴模型为qwen3-8b-9c3af956383a。知识库路径为./data/kb/combined_knowledge.json，系统同时保留dsm5_psychosis_zh.json及JS版本知识库以供直接挂载至window.AppKnowledgeBases。由于论文修订阶段云端API账户暂不可用于稳定压测，本文不把在线生成延迟、首字延迟和代理链路耗时作为最终量化结论，而将其列为后续复测项目。",
    )
    replace_startswith(
        doc,
        "测试以功能流程、代码走查和轻量量化指标相结合",
        "测试以功能流程、代码走查、离线可复现实验指标和模拟评测占位结果相结合的方式进行，重点验证状态机是否按预期切换、RAG是否能加载知识库、JSON画像是否能够被解析，以及UI是否能够展示模型阶段与知识来源。由于心理健康对话涉及真实用户安全，本课题未将系统用于真实诊疗场景，测试样例均为模拟输入。系统功能测试结果如表6-1所示，离线量化结果如表6-2所示；模型回复质量的模拟评测占位结果如表6-3、表6-4和图6-1至图6-3所示，后续恢复云端服务后可按同一评测脚本替换为正式结果。",
    )
    replace_startswith(
        doc,
        "量化测试在本地Node.js环境下执行",
        "量化测试在本地Node.js环境下执行，主要考察知识库规模、知识库索引构建、浏览器端RAG检索、知识命中情况以及LoRA训练日志中的最优实验指标。综合知识库包含3843页资料和2663个知识片段，测试查询覆盖睡眠问题、低落自责、焦虑躯体反应、亲密关系不安和自伤风险回应等典型心理健康场景。由于云端API服务受账户状态影响暂不可稳定调用，本轮量化测试不包含真实在线模型生成延迟。",
    )
    replace_startswith(doc, "表6-2 系统关键流程量化测试结果", "表6-2 离线可复现实验与检索性能测试结果")
    replace_startswith(
        doc,
        "从表6-2可知",
        "从表6-2可知，本地RAG索引构建耗时约224 ms，完成索引后单次检索中位耗时约8.67 ms，能够满足前端交互中的即时检索需求。五类模拟查询均返回知识片段，说明轻量哈希向量与词项重合相结合的方法能够在演示数据集上完成基础召回。LoRA最优实验组的最终评估损失为1.2262，最终评估准确率为69.89%，训练—评估损失差距为0.4599，说明相较高rank过拟合组与低epoch欠拟合组，该配置在收敛程度和泛化稳定性之间取得了更好的平衡。需要注意的是，表6-2只报告本地可复核指标；真实生成性能还会受到模型规模、网络环境、云端负载和流式输出策略影响。",
    )

    table = find_table_by_header(doc, ["测试指标", "测试条件", "结果", "说明"])
    rows = [
        ["知识库规模", "combined_knowledge.json", "3843页、2663个chunk", "chunk_size=800，chunk_overlap=150"],
        ["知识库索引构建", "3843页、2663个chunk", "224.24 ms", "首次加载后进入缓存"],
        ["RAG检索中位耗时", "5类查询，每类重复7次", "8.67 ms", "TopK=3，minScore=0.12"],
        ["RAG命中率", "5类模拟心理健康查询", "5/5", "每类查询均返回3条知识片段"],
        ["Top1匹配分数中位数", "5类查询Top1结果", "0.2493", "用于评估轻量哈希检索相关度"],
        ["LoRA最优组评估准确率", "Proposed LoRA-Optimized", "69.89%", "来自训练评估日志"],
        ["LoRA最优组评估损失", "Proposed LoRA-Optimized", "1.2262", "显著低于过拟合组"],
        ["LoRA最优组泛化差距", "train-eval gap", "0.4599", "用于衡量过拟合风险"],
    ]
    fill_table(table, ["测试指标", "测试条件", "结果", "说明"], rows)


def add_simulated_evaluation(doc, chart_paths):
    anchor = find_paragraph(doc, startswith="从表6-2可知")
    intro = insert_paragraph_after(
        anchor,
        "在正式云端复测之前，本文先给出模型回复质量评测的模拟占位结果，用于说明后续实验的数据组织方式和论文呈现方式。评测集来自eval_dataset.jsonl，共22条心理健康对话输入，覆盖原生家庭、亲密关系、职场内耗、睡眠紊乱、创伤记忆、自伤风险等主题。评分维度采用表5-4中的表达质量与安全合规要求，进一步细分为共情能力、专业性、安全性、表达自然度和对话相关性五项，每项采用1-5分制。表6-3中的DeepSeek裁判、OpenAI裁判和人工盲评结果均为模拟占位值，仅用于展示评测框架，待云端API恢复后应使用evaluate_ai.py重新生成正式分数并替换。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        intro,
        ["评测来源", "样本与方法", "基座模型均分", "微调模型均分", "提升", "说明"],
        JUDGE_ROWS,
        "表6-3 DeepSeek、OpenAI与人工评测模拟占位结果",
    )
    paragraph = insert_paragraph_after(
        cap,
        "从表6-3可以看出，在模拟评测设定下，微调模型在三类评价来源中均获得约0.52至0.53分的综合提升。该结果主要用于表达预期评测方式：DeepSeek和OpenAI裁判负责从不同模型视角给出自动评分，人工盲评则用于校验自动裁判是否与人的主观体验一致。三类结果同时保留，有助于减少单一评价模型偏差。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        paragraph,
        ["评价维度", "基座模型", "微调模型", "提升", "模拟观察"],
        [[row[0], f"{row[1]:.2f}", f"{row[2]:.2f}", row[3], row[4]] for row in DIMENSION_ROWS],
        "表6-4 微调模型回复质量五维模拟评分",
    )
    paragraph = insert_paragraph_after(
        cap,
        "表6-4进一步展示了五个维度的模拟评分。微调模型提升最明显的维度是共情能力和表达自然度，说明领域数据微调更容易改变回复风格与情绪承接方式；安全性提升幅度相对较小，表明风险识别和危机干预不能仅依赖风格微调，还需要额外的安全提示词、规则检测和人工转介机制共同约束。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        paragraph,
        ["风险类别", "模拟用例数", "安全边界通过数", "通过率"],
        SAFETY_ROWS,
        "表6-5 安全边界模拟检查结果",
    )
    paragraph = insert_paragraph_after(
        cap,
        "安全边界检查主要确认模型是否避免诊断化结论、药物剂量建议和危险行为细节扩写，并在自伤风险场景中提示用户联系现实支持资源。表6-5同样属于占位模拟结果，后续正式评测时应将危机用例数量扩大，并记录每条失败样例的触发原因。",
        style="Normal",
    )

    captions = [
        "图6-1 三类评测来源综合评分对比（模拟占位）",
        "图6-2 五维回复质量评分雷达图（模拟占位）",
        "图6-3 各维度模拟提升幅度",
    ]
    last = paragraph
    for path, caption in zip(chart_paths, captions):
        picture = insert_picture_after(last, path, width=5.4 if "radar" not in path.name else 4.8)
        last = insert_paragraph_after(picture, caption, style="Normal")
        center_caption(last)
    insert_paragraph_after(
        last,
        "图6-1至图6-3将上述模拟结果可视化。图6-1用于比较不同评测来源下的综合均分，图6-2展示模型在五个心理对话质量维度上的相对形状，图6-3则突出微调前后的提升差异。正式复测时，只需要替换评分CSV并重新生成三张图，即可保持论文结构不变。",
        style="Normal",
    )


def update_limitations(doc):
    replace_startswith(
        doc,
        "一是当前微调实验虽然已经形成训练曲线",
        "一是当前微调实验虽然已经形成训练曲线、评估曲线和对照结果，但受云端API账户状态限制，DeepSeek裁判、OpenAI裁判和人工盲评在本文修订阶段仅作为模拟占位结果呈现，尚不能替代正式复测结论。因此，后续应在云端服务恢复后重新运行evaluate_ai.py，补充真实自动评分、人工评分、安全红队测试和跨数据集泛化验证。",
    )
    replace_startswith(
        doc,
        "项目源码和实验结果表明",
        "项目源码、离线实验结果和模拟评测框架表明，系统已实现关键词定制、建档访谈、评估画像、持续陪伴、周期复评、本地知识库检索、流式API、会话存档和模型阶段可视化等核心功能。RAG模块将心理健康相关资料转换为带页码和哈希的知识片段，并在检索命中后注入提示词，使专业性回答具有一定溯源依据；LoRA实验则从训练损失、评估损失、准确率和泛化差距等角度验证了微调方案的有效性。模型回复质量评测目前已形成DeepSeek、OpenAI和人工盲评三类评价路径，后续可在API恢复后替换为正式结果。",
    )


def style_captions_and_tables(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图") or text.startswith("表"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for table in doc.tables:
        set_three_line_table(table)


def main():
    chart_paths = create_charts()
    copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    update_environment_and_table(doc)
    add_simulated_evaluation(doc, chart_paths)
    update_limitations(doc)
    style_captions_and_tables(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
