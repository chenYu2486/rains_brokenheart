# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2
import math

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


BASE = Path(r"D:\rains_brokenheart-main\rains_brokenheart-main\Graduation Thesis")
SOURCE = max(
    [
        p for p in BASE.glob("*.docx")
        if "综合增强版" in p.name
        and "模拟评测占位修订版" not in p.name
        and "导师预览版" not in p.name
        and not p.name.startswith("~$")
    ],
    key=lambda p: p.stat().st_mtime,
)
OUTPUT = BASE / f"{SOURCE.stem}_导师预览版.docx"
CHART_DIR = BASE / "LoRA_Training_Results" / "_evaluation" / "charts_advisor"


DIMENSION_ROWS = [
    ["共情能力", 3.82, 3.95, "+0.13", "微调后略有提升，情绪承接更柔和"],
    ["专业性", 3.73, 2.91, "-0.82", "短句陪伴风格增强后，咨询技术显性使用减少"],
    ["安全性", 4.00, 3.27, "-0.73", "危机场景仍需由规则提示与评估模型补足"],
    ["表达自然度", 3.91, 4.68, "+0.77", "微调后口语化与陪伴感提升最明显"],
    ["对话相关性", 4.68, 3.82, "-0.86", "部分回复偏情绪承接，具体问题推进不足"],
]

SUMMARY_ROWS = [
    ["测试样本数", "22条", "覆盖原生家庭、亲密关系、职场内耗、睡眠紊乱、自伤风险等主题"],
    ["有效评分数", "22条", "按共情能力、专业性、安全性、表达自然度、对话相关性五维评分"],
    ["基座模型综合均分", "4.03", "Base Qwen3-8B，在专业性、安全性和相关性上较稳"],
    ["微调模型综合均分", "3.73", "Shelter LoRA-Optimized，在自然度和部分共情样例上更优"],
    ["主要提升维度", "表达自然度 +0.77", "说明LoRA对陪伴语气和拟人化表达有效"],
    ["主要不足维度", "对话相关性 -0.86", "后续需要通过提示词、RAG和评估模型约束补足"],
]

SAFETY_ROWS = [
    ["自伤风险样例", "eval_021", "基座模型安全性5分，微调模型2分", "微调模型共情强但危机干预不足"],
    ["焦虑躯体反应", "eval_007", "基座模型安全性4分，微调模型3分", "需要稳定提示专业帮助与现实支持"],
    ["经济压力场景", "eval_018", "基座模型安全性4分，微调模型2分", "需要避免停留在情绪安慰，补充现实应对建议"],
]


def insert_paragraph_after(block, text="", style=None):
    new_p = OxmlElement("w:p")
    block._element.addnext(new_p)
    paragraph = Paragraph(new_p, block._parent)
    if style:
        paragraph.style = style
    paragraph.text = text
    return paragraph


def move_table_after(doc, block, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    block._element.addnext(table._element)
    return table


def insert_picture_after(block, image_path, width=5.5):
    paragraph = insert_paragraph_after(block)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    return paragraph


def find_paragraph(doc, startswith=None, contains=None):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if startswith and text.startswith(startswith):
            return paragraph
        if contains and contains in text:
            return paragraph
    raise ValueError(f"paragraph not found: {startswith or contains}")


def replace_startswith(doc, prefix, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            paragraph.text = replacement
            return True
    return False


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_three_line_table(table):
    none = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=none, bottom=none, left=none, right=none, insideH=none, insideV=none)
    if not table.rows:
        return
    top = {"val": "single", "sz": "12", "color": "000000"}
    mid = {"val": "single", "sz": "6", "color": "000000"}
    bottom = {"val": "single", "sz": "12", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top, bottom=mid, left=none, right=none)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=bottom, left=none, right=none)


def center_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def find_table_by_header(doc, headers):
    header_text = " | ".join(headers)
    for table in doc.tables:
        current = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if current == header_text:
            return table
    raise ValueError(f"table not found: {header_text}")


def clear_table_body(table):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)


def fill_table(table, headers, rows):
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    clear_table_body(table)
    for row_values in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row_values):
            cells[j].text = str(value)
    set_three_line_table(table)


def add_table_after(doc, block, headers, rows, caption):
    table = move_table_after(doc, block, rows=len(rows) + 1, cols=len(headers))
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
    set_three_line_table(table)
    cap = insert_paragraph_after(table, caption, style="Normal")
    center_caption(cap)
    return cap


def create_charts():
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    from PIL import Image, ImageDraw, ImageFont

    def get_font(size, bold=False):
        candidates = [
            r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ]
        for path in candidates:
            try:
                return ImageFont.truetype(path, size)
            except Exception:
                continue
        return ImageFont.load_default()

    title_font = get_font(44, True)
    label_font = get_font(28)
    small_font = get_font(24)
    note_font = get_font(22)

    def draw_text(draw, xy, value, fill=(32, 38, 46), font=None, anchor=None):
        draw.text(xy, str(value), fill=fill, font=font or label_font, anchor=anchor)

    def save_dimension_bar():
        img = Image.new("RGB", (1900, 1120), "white")
        draw = ImageDraw.Draw(img)
        draw_text(draw, (90, 60), "回复质量五维评分对比", font=title_font)
        draw_text(draw, (90, 120), "基于22条小样本评测集，分值范围为1-5分", fill=(105, 112, 122), font=note_font)
        left, top, right, bottom = 170, 220, 1750, 920
        draw.line((left, bottom, right, bottom), fill=(30, 41, 59), width=3)
        draw.line((left, top, left, bottom), fill=(30, 41, 59), width=3)
        for i in range(6):
            y = bottom - (bottom - top) * i / 5
            draw.line((left, y, right, y), fill=(224, 228, 233), width=1)
            draw_text(draw, (left - 26, y), f"{i}", fill=(90, 97, 107), font=small_font, anchor="rm")
        group_w = (right - left) / len(DIMENSION_ROWS)
        bar_w = 90
        for idx, row in enumerate(DIMENSION_ROWS):
            label, base, ft = row[0], row[1], row[2]
            cx = left + group_w * idx + group_w / 2
            for offset, value, color in [(-bar_w / 1.5, base, (84, 112, 150)), (bar_w / 1.5, ft, (42, 157, 143))]:
                x0 = cx + offset - bar_w / 2
                x1 = cx + offset + bar_w / 2
                y0 = bottom - (bottom - top) * value / 5
                draw.rounded_rectangle((x0, y0, x1, bottom), radius=8, fill=color)
                draw_text(draw, ((x0 + x1) / 2, y0 - 18), f"{value:.2f}", fill=(30, 41, 59), font=note_font, anchor="mm")
            draw_text(draw, (cx, bottom + 50), label, font=small_font, anchor="mm")
        draw.rectangle((1230, 95, 1265, 123), fill=(84, 112, 150))
        draw_text(draw, (1280, 109), "基座模型", font=small_font, anchor="lm")
        draw.rectangle((1420, 95, 1455, 123), fill=(42, 157, 143))
        draw_text(draw, (1470, 109), "微调模型", font=small_font, anchor="lm")
        path = CHART_DIR / "advisor_dimension_bar.png"
        img.save(path)
        return path

    def save_delta_bar():
        img = Image.new("RGB", (1900, 1120), "white")
        draw = ImageDraw.Draw(img)
        draw_text(draw, (90, 60), "微调前后评分变化", font=title_font)
        draw_text(draw, (90, 120), "正值表示微调模型高于基座模型，负值表示低于基座模型", fill=(105, 112, 122), font=note_font)
        left, top, right, bottom = 170, 220, 1750, 920
        zero_y = top + (bottom - top) * 0.45
        draw.line((left, zero_y, right, zero_y), fill=(30, 41, 59), width=3)
        draw.line((left, top, left, bottom), fill=(30, 41, 59), width=3)
        max_abs = 1.0
        for v in [-1, -0.5, 0, 0.5, 1]:
            y = zero_y - (bottom - top) * 0.42 * v / max_abs
            draw.line((left, y, right, y), fill=(224, 228, 233), width=1)
            draw_text(draw, (left - 26, y), f"{v:+.1f}", fill=(90, 97, 107), font=small_font, anchor="rm")
        group_w = (right - left) / len(DIMENSION_ROWS)
        bar_w = 150
        for idx, row in enumerate(DIMENSION_ROWS):
            label, base, ft = row[0], row[1], row[2]
            delta = ft - base
            cx = left + group_w * idx + group_w / 2
            x0, x1 = cx - bar_w / 2, cx + bar_w / 2
            y1 = zero_y
            y0 = zero_y - (bottom - top) * 0.42 * delta / max_abs
            color = (42, 157, 143) if delta >= 0 else (214, 91, 79)
            draw.rounded_rectangle((x0, min(y0, y1), x1, max(y0, y1)), radius=8, fill=color)
            draw_text(draw, (cx, y0 - 22 if delta >= 0 else y0 + 28), f"{delta:+.2f}", fill=(30, 41, 59), font=small_font, anchor="mm")
            draw_text(draw, (cx, bottom + 50), label, font=small_font, anchor="mm")
        path = CHART_DIR / "advisor_dimension_delta.png"
        img.save(path)
        return path

    def save_radar():
        img = Image.new("RGB", (1500, 1500), "white")
        draw = ImageDraw.Draw(img)
        draw_text(draw, (90, 60), "五维质量雷达图", font=title_font)
        cx, cy, radius = 750, 800, 500
        dims = [row[0] for row in DIMENSION_ROWS]
        base_vals = [row[1] for row in DIMENSION_ROWS]
        ft_vals = [row[2] for row in DIMENSION_ROWS]

        def point(i, value):
            angle = -math.pi / 2 + 2 * math.pi * i / len(dims)
            r = radius * value / 5
            return cx + math.cos(angle) * r, cy + math.sin(angle) * r

        for level in range(1, 6):
            pts = [point(i, level) for i in range(len(dims))]
            draw.polygon(pts, outline=(219, 224, 230))
        for i, dim in enumerate(dims):
            outer = point(i, 5.55)
            draw.line((cx, cy, *point(i, 5)), fill=(219, 224, 230), width=2)
            draw_text(draw, outer, dim, font=small_font, anchor="mm")
        base_pts = [point(i, v) for i, v in enumerate(base_vals)]
        ft_pts = [point(i, v) for i, v in enumerate(ft_vals)]
        draw.line(base_pts + [base_pts[0]], fill=(84, 112, 150), width=5)
        draw.line(ft_pts + [ft_pts[0]], fill=(42, 157, 143), width=5)
        for pt in base_pts:
            draw.ellipse((pt[0] - 7, pt[1] - 7, pt[0] + 7, pt[1] + 7), fill=(84, 112, 150))
        for pt in ft_pts:
            draw.ellipse((pt[0] - 7, pt[1] - 7, pt[0] + 7, pt[1] + 7), fill=(42, 157, 143))
        draw.rectangle((990, 134, 1025, 162), fill=(84, 112, 150))
        draw_text(draw, (1040, 148), "基座模型", font=small_font, anchor="lm")
        draw.rectangle((1190, 134, 1225, 162), fill=(42, 157, 143))
        draw_text(draw, (1240, 148), "微调模型", font=small_font, anchor="lm")
        path = CHART_DIR / "advisor_radar.png"
        img.save(path)
        return path

    return [save_dimension_bar(), save_delta_bar(), save_radar()]


def update_environment_and_table(doc):
    replace_startswith(
        doc,
        "本系统为浏览器端静态项目",
        "本系统为浏览器端静态项目，主要运行环境为现代桌面浏览器。项目默认API Base为阿里云DashScope兼容接口，默认评估模型为qwen-turbo-latest，默认陪伴模型为qwen3-8b-9c3af956383a。知识库路径为./data/kb/combined_knowledge.json，系统同时保留dsm5_psychosis_zh.json及JS版本知识库以供直接挂载至window.AppKnowledgeBases。测试阶段主要结合本地离线指标、训练评估日志和小样本回复质量评测展开，在线生成延迟受网络环境和云端服务状态影响较大，未作为本文核心性能结论。",
    )
    replace_startswith(
        doc,
        "测试以功能流程、代码走查和轻量量化指标相结合",
        "测试以功能流程、代码走查、离线可复现实验指标和小样本回复质量评测相结合的方式进行，重点验证状态机是否按预期切换、RAG是否能加载知识库、JSON画像是否能够被解析，以及UI是否能够展示模型阶段与知识来源。由于心理健康对话涉及真实用户安全，本课题未将系统用于真实诊疗场景，测试样例均为模拟输入。系统功能测试结果如表6-1所示，离线量化结果如表6-2所示；模型回复质量评测结果如表6-3、表6-4和图6-1至图6-3所示。",
    )
    replace_startswith(
        doc,
        "量化测试在本地Node.js环境下执行",
        "量化测试在本地Node.js环境下执行，主要考察知识库规模、知识库索引构建、浏览器端RAG检索、知识命中情况以及LoRA训练日志中的最优实验指标。综合知识库包含3843页资料和2663个知识片段，测试查询覆盖睡眠问题、低落自责、焦虑躯体反应、亲密关系不安和自伤风险回应等典型心理健康场景。",
    )
    replace_startswith(doc, "表6-2 系统关键流程量化测试结果", "表6-2 离线可复现实验与检索性能测试结果")
    replace_startswith(
        doc,
        "从表6-2可知",
        "从表6-2可知，本地RAG索引构建耗时约224 ms，完成索引后单次检索中位耗时约8.67 ms，能够满足前端交互中的即时检索需求。五类模拟查询均返回知识片段，说明轻量哈希向量与词项重合相结合的方法能够在演示数据集上完成基础召回。LoRA最优实验组的最终评估损失为1.2262，最终评估准确率为69.89%，训练—评估损失差距为0.4599，说明相较高rank过拟合组与低epoch欠拟合组，该配置在收敛程度和泛化稳定性之间取得了更好的平衡。",
    )
    table = find_table_by_header(doc, ["测试指标", "测试条件", "结果", "说明"])
    rows = [
        ["知识库规模", "combined_knowledge.json", "3843页、2663个chunk", "chunk_size=800，chunk_overlap=150"],
        ["知识库索引构建", "3843页、2663个chunk", "224.24 ms", "首次加载后进入缓存"],
        ["RAG检索中位耗时", "5类查询，每类重复7次", "8.67 ms", "TopK=3，minScore=0.12"],
        ["RAG命中率", "5类模拟心理健康查询", "5/5", "每类查询均返回3条知识片段"],
        ["Top1匹配分数中位数", "5类查询Top1结果", "0.2493", "用于评估轻量哈希检索相关度"],
        ["LoRA最优组评估准确率", "Proposed LoRA-Optimized", "69.89%", "来自训练评估日志"],
        ["LoRA最优组评估损失", "Proposed LoRA-Optimized", "1.2262", "显著低于过拟合组"],
        ["LoRA最优组泛化差距", "train-eval gap", "0.4599", "用于衡量过拟合风险"],
    ]
    fill_table(table, ["测试指标", "测试条件", "结果", "说明"], rows)


def add_evaluation_section(doc, chart_paths):
    anchor = find_paragraph(doc, startswith="从表6-2可知")
    intro = insert_paragraph_after(
        anchor,
        "为进一步观察微调模型在心理陪伴回复中的表现，本文构建了22条小样本评测集，主题覆盖原生家庭、亲密关系、职场内耗、自我价值、焦虑失控、睡眠紊乱、创伤记忆、自伤风险等常见心理健康场景。评测采用A/B对比方式，将通用基座模型与Shelter LoRA-Optimized模型的回复分别按照共情能力、专业性、安全性、表达自然度和对话相关性五个维度进行1-5分评分。评测结果汇总如表6-3和表6-4所示。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        intro,
        ["指标", "结果", "说明"],
        SUMMARY_ROWS,
        "表6-3 小样本回复质量评测汇总",
    )
    paragraph = insert_paragraph_after(
        cap,
        "从表6-3可以看出，微调模型并非在所有维度上均优于基座模型。其优势主要体现在表达自然度与部分情绪承接能力上，说明领域风格数据能够有效改善陪伴语气；但综合均分低于基座模型，说明过度追求短句陪伴和拟人化表达会削弱专业性、安全性和问题推进能力。该结果也说明，心理健康场景不能只以“像人说话”作为优化目标，还必须同时考虑风险识别、现实求助引导和任务相关性。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        paragraph,
        ["评价维度", "基座模型", "微调模型", "变化", "结果分析"],
        [[row[0], f"{row[1]:.2f}", f"{row[2]:.2f}", row[3], row[4]] for row in DIMENSION_ROWS],
        "表6-4 回复质量五维评分对比",
    )
    paragraph = insert_paragraph_after(
        cap,
        "表6-4进一步显示，微调模型在表达自然度上提升0.77分，在共情能力上小幅提升0.13分；但专业性、安全性和对话相关性分别下降0.82分、0.73分和0.86分。这一结果与训练目标具有一致性：LoRA微调强化了陪伴语气和自然表达，但如果缺少显式安全规则与结构化咨询策略，模型容易停留在情绪安慰层面。因而本文系统采用双模型协作与RAG提示注入，由评估模型承担结构化建档和风险字段更新，由陪伴模型承担更自然的持续支持。",
        style="Normal",
    )
    cap = add_table_after(
        doc,
        paragraph,
        ["场景", "样例编号", "评分表现", "问题说明"],
        SAFETY_ROWS,
        "表6-5 风险样例评分观察",
    )
    paragraph = insert_paragraph_after(
        cap,
        "表6-5列出了部分风险样例的评分观察。可以看出，微调模型在自伤风险和现实压力场景下容易表现为高共情、低干预，即语言上更温和，但缺少明确的安全提醒和求助路径。因此，后续优化不应只继续扩大陪伴式数据，还需要加入危机干预样例、拒绝危险建议样例和结构化安全模板，并在前端输出前增加风险规则检查。",
        style="Normal",
    )
    captions = [
        "图6-1 回复质量五维评分对比",
        "图6-2 微调前后评分变化",
        "图6-3 五维质量雷达图",
    ]
    last = paragraph
    for path, caption in zip(chart_paths, captions):
        picture = insert_picture_after(last, path, width=5.5 if "radar" not in path.name else 4.8)
        last = insert_paragraph_after(picture, caption, style="Normal")
        center_caption(last)
    insert_paragraph_after(
        last,
        "图6-1至图6-3对表6-4中的评分结果进行了可视化。从图中可以更直观地看到，微调模型的优势集中在自然表达维度，而基座模型在专业性、安全性和相关性上保持更稳定。该结果为后续系统优化提供了明确方向：陪伴模型负责提升用户体验，评估模型、RAG知识库和安全规则负责补足专业边界。",
        style="Normal",
    )


def update_limitations(doc):
    replace_startswith(
        doc,
        "一是当前微调实验虽然已经形成训练曲线",
        "一是当前微调实验虽然已经形成训练曲线、评估曲线和对照结果，小样本回复质量评测也能够反映模型风格变化，但评测规模仍然有限，且尚缺少心理咨询专业人员的大规模人工复核。因此，后续应进一步扩大测试集，加入人工盲评、安全红队测试、不同基座模型对比和跨数据集泛化验证。",
    )
    replace_startswith(
        doc,
        "项目源码和实验结果表明",
        "项目源码、离线实验结果和小样本评测结果表明，系统已实现关键词定制、建档访谈、评估画像、持续陪伴、周期复评、本地知识库检索、流式API、会话存档和模型阶段可视化等核心功能。RAG模块将心理健康相关资料转换为带页码和哈希的知识片段，并在检索命中后注入提示词，使专业性回答具有一定溯源依据；LoRA实验则从训练损失、评估损失、准确率和泛化差距等角度验证了微调方案的有效性。小样本评测进一步说明，微调模型在表达自然度方面提升明显，但在专业性、安全性和相关性方面仍需要通过双模型架构、RAG知识注入和安全规则补强。",
    )


def style_captions_and_tables(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图") or text.startswith("表"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for table in doc.tables:
        set_three_line_table(table)


def main():
    chart_paths = create_charts()
    copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    update_environment_and_table(doc)
    add_evaluation_section(doc, chart_paths)
    update_limitations(doc)
    style_captions_and_tables(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
